import tkinter as tk
import tkinter.messagebox
from tkinter import filedialog
from tkinter import messagebox
# from tkinter import ttk
from json_to_pandas import Export_test_rail
import pandas as pd
import yaml
from docx.document import Document
from docx import Document
import os

class Main_app(tk.Tk):
    def __init__(self):
        super().__init__()
        LBL_PAD = 10

        # row item
        EXP_TYPE = 0
        PROJECT = 1
        T_RUN = 2
        TEMP = 4
        OUT = 5
        SEC_TBL = 7
        EXP_CLOSE = SEC_TBL+5

        if not os.path.isfile("./config.yml"):

            with open("config.yml", 'w') as f:
                f.writelines("---\n")
                f.writelines("project id: 2\n")
                f.writelines("test report: False\n")
                f.writelines("table mapping: {}\n")
                f.writelines("test report run id:\n")
                f.writelines("template path:\n")
                f.writelines("output doc name:\n")




        with open("config.yml") as f:
            config = yaml.safe_load(f)  # load data in config.yml in a dict

        self.export_type_var = tk.BooleanVar()  # False --> Specification ; True --> Report
        self.export_type_var.set(config["test report"])
        self.word_table_id_var = tk.IntVar()
        self.test_rail_id_var = tk.IntVar()
        self.table_section_text = tk.StringVar()
        self.template_path_var = tk.StringVar()
        self.template_path_var.set(config["template path"])
        self.output_file_path_var = tk.StringVar()
        self.output_file_path_var.set(config["output doc name"])
        self.project_id_var = tk.IntVar()
        self.project_id_var.set(config["project id"])
        self.test_report_run_id_var = tk.IntVar()

        self.table_map = {}

        self.geometry("860x440")
        self.wm_minsize(width=860, height=440)

        self.section_frame = tk.LabelFrame(self, width=500, height=100, bd=3, relief="sunken")
        self.section_frame.grid(column=1, pady=0, row=SEC_TBL+4, columnspan=3, sticky=tk.W)
        self.section_frame.grid_propagate(False)

        self.export_type_spec = tk.Radiobutton(self, text="Specification", variable=self.export_type_var, value=False)
        self.export_type_spec.grid(column=1, row=EXP_TYPE, padx=0, pady=5, sticky='w')
        self.export_type_rep = tk.Radiobutton(self, text="Report", variable=self.export_type_var, value=True)
        self.export_type_rep.grid(column=1, row=EXP_TYPE, padx=100, pady=5, sticky='w')

        self.project_id_lbl = tk.Label(self, text="Project ID")
        self.project_id_lbl.grid(column=0, row=PROJECT, padx=LBL_PAD, sticky='e')
        self.project_id_entry = tk.Entry(self, width=5, textvariable=self.project_id_var)
        self.project_id_entry.grid(column=1, row=PROJECT, sticky='w')
        self.project_id_comment = tk.Label(self, text="2 - Introral Scanner, 5 - DWOS Lab Management")
        self.project_id_comment.grid(column=1, row=PROJECT, padx= 30, ipadx=8, sticky=tk.W)
        self.test_report_run_id_lbl = tk.Label(self, text="Test Run ID")
        self.test_report_run_id_lbl.grid(column=0, row=T_RUN, padx=LBL_PAD, sticky='e')
        self.test_report_run_id_entry = tk.Entry(self, textvariable=self.test_report_run_id_var)
        self.test_report_run_id_entry.grid(column=1, row=T_RUN, sticky=tk.W)
        self.test_report_run_id_comment = tk.Label(text="Optional, value use only for reporting")
        self.test_report_run_id_comment.grid(column=1, row=T_RUN, padx= 100, ipadx=8, sticky=tk.W)

        self.space = tk.Label(text="", font=("Arial", 4))
        self.space.grid(column=0, row=3)

        self.template_lbl = tk.Label(self, text="Template path")
        self.template_lbl.grid(column=0, row=TEMP, padx=LBL_PAD, pady=5, sticky=tk.SE)
        self.template_entry = tk.Entry(self, width=110, textvariable=self.template_path_var)
        self.template_entry.grid(column=1, row=TEMP, pady=5, ipady=3, sticky=tk.S)
        self.template_browse = tk.Button(self, text="Browse",
                                         command=lambda: self.browse(self.template_path_var, mode="open",
                                                                     initialdir=self.output_file_path_var.get()))
        self.template_browse.grid(column=3, row=TEMP, padx=10, sticky='w')
        self.output_file_lbl = tk.Label(self, text="Output file path")
        self.output_file_lbl.grid(column=0, row=OUT, padx=LBL_PAD, sticky='e')
        self.output_file_entry = tk.Entry(self, width=110, textvariable=self.output_file_path_var)
        self.output_file_entry.grid(column=1, row=OUT, ipady=3, sticky='w')
        self.output_file_browse = tk.Button(self, text="Browse",
                                            command=lambda: self.browse(self.output_file_path_var, mode="save",
                                                                        initialdir=self.output_file_path_var.get(),
                                                                        filetypes=[("Word Doc", ".docx")],
                                                                        defaultextension=".docx"))
        self.output_file_browse.grid(column=3, row=OUT, padx=10)

        self.space = tk.Label(text="", font=("Arial", 4))
        self.space.grid(column=0, row=6)

        self.test_rail_id_lbl = tk.Label(self, text="Test Rail section ID")
        self.test_rail_id_lbl.grid(column=0, row=SEC_TBL, padx=LBL_PAD, sticky='e')
        self.test_rail_id_entry = tk.Entry(self)
        self.test_rail_id_entry.grid(column=1, row=SEC_TBL, sticky='w')
        self.word_table_id_lbl = tk.Label(self, text="Word Table ID")
        self.word_table_id_lbl.grid(column=0, row=SEC_TBL+1, padx=LBL_PAD, sticky='e')
        self.word_table_id_entry = tk.Entry(self)
        self.word_table_id_entry.grid(column=1, row=SEC_TBL+1, sticky='w')
        self.add_section_btn = tk.Button(self, text="Add", command=self.add_table_section_element)
        self.add_section_btn.grid(column=1, row=SEC_TBL+2, padx=0, pady=5, sticky='w')
        self.clear_section_btn = tk.Button(self, text="Clear", command=self.clear_section_id_element)
        self.clear_section_btn.grid(column=1, row=SEC_TBL+2, padx=50, pady=5, sticky='w')

        self.frame_label = tk.Label(text="List of exported section (section : word table)")
        self.frame_label.grid(column=1, row=SEC_TBL+3, sticky=tk.W)

        self.table_section = tk.Label(self.section_frame, textvariable=self.table_section_text, justify=tk.LEFT)
        self.table_section.grid(column=0, row=0, padx=0, sticky=tk.W)

        self.validate = tk.Button(self, text="Export", command=self.validate)
        self.validate.grid(column=1, row=EXP_CLOSE, padx=0, pady=30, ipadx=50, sticky='w')
        self.close = tk.Button(self, text="Close", command=self.close)
        self.close.grid(column=1, row=EXP_CLOSE, padx=170, pady=30, sticky='w')


        cc = {"project id": self.project_id_var,
              "test report run id": self.test_report_run_id_var,
              "template path": self.template_path_var,
              "output doc name": self.output_file_path_var,
              "test report": self.export_type_var
              }

        try:
            with open("config.yml") as f:
                self.existing_config = yaml.safe_load(f)  # load data in config.yml in a dict
                print(self.existing_config)
                if len(self.existing_config) > 0:
                    for key, value in cc.items():
                        value.set(self.existing_config[key])
                        # self.project_id_var.set(self.existing_config["project id"])
        except Exception as e:
            print(e)

    def validate(self):
        print(self.export_type_var.get())
        print(self.template_path_var.get())
        print(self.output_file_path_var.get())
        print(self.table_map)
        if len(self.table_map) > 0:
            with open("config.yml", mode='w') as f:
                f.writelines("---\n")
                f.writelines("project id: " + str(self.project_id_var.get()) + "\n")
                f.writelines("test report: " + str(self.export_type_var.get()) + "\n")
                f.writelines(self.create_yml_dict(self.table_map, "table mapping") + "\n")
                f.writelines("test report run id: " + str(self.test_report_run_id_var.get()) + "\n")
                f.writelines("template path: " + self.template_path_var.get() + "\n")
                f.writelines("output doc name: " + self.output_file_path_var.get() + "\n")

            self.ex = Export_test_rail()
            user_initial = self.ex.create_user_initial()

            if self.export_type_var.get():  # it's a Report
                self.create_report()
            else:  # it's a specification
                self.create_specification()
        else:
            tkinter.messagebox.showerror(title="Export section list empty",
                                         message="Please add section to export by clicking on ADD button")

    def close(self):
        self.destroy()

    def create_yml_dict(self, data, tag_name):
        data_string = tag_name + ": {"
        for key, value in data.items():
            data_string += str(key) + ": " + str(value) + ", "
        data_string += "}"

        return data_string

    def add_table_section_element(self):
        text = ""
        try:
            self.table_map[self.test_rail_id_entry.get()] = str(int(self.word_table_id_entry.get()) - 1)
        except ValueError:
            tk.messagebox.showerror(title="Error", message="Cannot add empty field")
        row_element = 0
        for key, value in self.table_map.items():
            if row_element < 5:
                text = text + str(key) + " : " + str(int(value)+1) + ", "
                row_element += 1
            else:
                text = text + str(key) + " : " + str(int(value)+1) + "\n"
                row_element = 0
        self.table_section_text.set(text)

    def clear_section_id_element(self):
        self.table_map = {}
        self.table_section_text.set("")
        self.update()

    def browse(self, var, **kwargs):
        if kwargs["mode"] == "open":
            path = filedialog.askopenfilename(initialdir=kwargs["initialdir"])
        else:
            path = filedialog.asksaveasfilename(initialdir=kwargs["initialdir"], filetypes=kwargs["filetypes"],
                                                defaultextension=kwargs["defaultextension"])
        print(path)
        var.set(path)

    def browse_folder(self, var):
        path = filedialog.askdirectory()
        print(path)
        var.set(path)

    def create_report(self):
        print("Process test Report")
        document = Document(self.ex.config["template path"])
        for tr_section_id, table_id in self.ex.config["table mapping"].items():
            tc_list = self.ex.get_tc_list_from_section(tr_section_id)
            self.ex.test_run_step_counter = 0
            df_orginal = None
            for test_case_id in tc_list:
                url = "https://testrail.dwos.com/index.php?/api/v2/get_case/" + str(test_case_id)
                self.ex.tc_js = self.ex.my_http_request(url, self.ex.USER, self.ex.PASSWORD)

                url = "https://testrail.dwos.com//index.php?/api/v2/get_results_for_case/" + \
                      str(self.ex.config["test report run id"]) + "/" + str(test_case_id)
                self.ex.js = self.ex.my_http_request(url, self.ex.USER, self.ex.PASSWORD)
                try:
                    self.ex.js[0]
                    js = self.ex.get_last_valid_result(self.ex.js)
                    df = self.ex.create_df_from_json(js, ["Step", "Result Description", "Result", "Date", "Initial"],
                                                     [self.ex.tr_step_num, self.ex.tr_result_description,
                                                      self.ex.tr_result, self.ex.tr_date,
                                                      self.ex.tr_initial], test_case_id)
                    df_orginal = pd.concat([df_orginal, df], ignore_index=True)
                except:
                    continue
        # self.ex.write_to_doc_table(df_orginal, table_id, self.ex.config["template path"],
        #                            self.ex.config["output doc name"])
            self.ex.write_to_doc_table(df_orginal, table_id, document, self.ex.config["output doc name"])
        tk.messagebox.showinfo(title="Done", message="Export Completed")

    def create_specification(self):
        document = Document(self.ex.config["template path"])
        for tr_section_id, table_id in self.ex.config["table mapping"].items():
            url = "https://testrail.dwos.com//index.php?/api/v2/get_cases/" + str(
                self.ex.config["project id"]) + "&section_id=" + str(tr_section_id)
            js = self.ex.my_http_request(url, self.ex.USER, self.ex.PASSWORD)
            df = self.ex.create_df_from_json(js, ["Step", "Description", "Requirement", "Expected Result",
                                                  "Test Method / Objective Evidence"],
                                             [self.ex.step_num, self.ex.tc_description, self.ex.tc_requirement,
                                              self.ex.tc_expected_result, self.ex.tc_test_method])
            # self.ex.write_to_doc_table(df, table_id, self.ex.config["template path"], self.ex.config["output doc name"])
            self.ex.write_to_doc_table(df, table_id, document, self.ex.config["output doc name"])
        tk.messagebox.showinfo(title="Done", message="Export Completed")


app = Main_app()
app.mainloop()
