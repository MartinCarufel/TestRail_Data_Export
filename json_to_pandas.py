import pandas as pd
import json
from docx.document import Document
from docx.shared import Pt
from docx import Document
import re
import yaml
import datetime
import requests


class Export_test_rail:

    def __init__(self):


        self.USER = 'martin.carufel@dental-wings.com'
        self.PASSWORD = '18,Mac&Amo'

        pd.set_option("display.max_columns", 5)
        self.status_text = {1:"Pass",
                       2:"Blocked",
                       3:"Untested",
                       4:"Retest",
                       5:"Fail",
                       7:"Pass w/Dev"}
        self.user_initial = self.create_user_initial()
        self.test_run_step_counter = 0
        self.js = None
        self.tc_js = None

        with open('config.yml') as f:
            self.config = yaml.safe_load(f)

    def create_user_initial(self):
        self.user_initial = {60: "ATR",
                        45: "AAR",
                        11: "ABN",
                        62: "BLT",
                        18: "CPT",
                        37: "CAA",
                        32: "CGS",
                        14: "DRR",
                        31: "DGU",
                        24: "JHE",
                        54: "KTU",
                        12: "MCE",
                        41: "MCL",
                        51: "NPN",
                        29: "PUR",
                        47: "PUR",
                        61: "PCR",
                        63: "SME",
                        59: "TLS",
                        950: "VDD"
        }



    def remove_picture_placeholder(self, input_text):
        regex_patern = '!\\[\\]\\(index\\.php\\?/attachments/get/\\d+\\) *\n*'
        output_text = re.subn(regex_patern, "", input_text)
        return output_text[0]

    def get_expected_result(self, data):
        """
        :param data: Is the content of the test case dict 'custom_steps_separated'
        :return: String of all expected results
        """
        expected_result = []
        for i in range(len(data)):
            if i == len(data) - 1:  # prevent to add empty line at the end of the table
                expected_result.append("Step {}: {}".format(str(i+1), data[i]["expected"]))
            else:
                expected_result.append("Step {}: {}\n".format(str(i + 1), data[i]["expected"]))

        return "\n".join(expected_result)

    def extract_all_step_result(self, custom_step_result):
        """
        :param custom_step_result: Provide the content of the "get_results_for_case["custom_step_results"]" of
        a valid result
        :return:
        """
        regex_patern = '!\\[\\]\\(index\\.php\\?/attachments/get/\\d+\\) *\n*'
        steps_results = []

        for actual_result_idx in range(len(custom_step_result)):
            if custom_step_result[actual_result_idx]["actual"] != "":
                steps_results.append("Step {}: ".format(actual_result_idx + 1))
                filtered_string = re.subn(regex_patern, "", self.remove_picture_placeholder(
                    custom_step_result[actual_result_idx]["actual"]))
                steps_results.append(filtered_string[0])
        return "\n".join(steps_results)


    def write_to_doc_table(self, df, table_id, template_doc, output_doc):
        d = Document(template_doc)
        df_nb_row, df_nb_col = df.shape
        for df_row_idx in range(df_nb_row):
            table_row = df_row_idx + 1
            d.tables[table_id].add_row()
            for col_idx in range(df_nb_col):
                # print("row: {}   col: {}".format(df_row_idx, col_idx))
                d.tables[table_id].cell(table_row, col_idx).text = str(df.loc[df_row_idx][col_idx])
        self.change_table_font(d.tables[table_id], 9)
        d.save(output_doc)

    def print_dataframe(self, df):
        col_names = df.columns
        print(col_names)
        nb_row, nb_col = df.shape
        for row_idx in range(nb_row):
            print("-------------------------------------------------------------")
            for col_idx in range(nb_col):
                print("{}: {}".format(col_names[col_idx], df.loc[row_idx][col_idx]))


    def create_df_from_json(self, json_obj, col_list, col_content, tc=None):
        df = pd.DataFrame(columns=col_list)
        for row_idx in range(len(json_obj)):
            row_data = []
            for col_idx in range(len(col_list)):
                row_data.append(col_content[col_idx](json_obj[row_idx], row_idx, tc))
            df.loc[row_idx] = row_data
        # print(df)
        return df


    def step_num(self, json_obj, row_idx, tc):
        return str(row_idx + 1)

    def tc_description(self, json_obj, row_idx, tc):
        return "Test Case: C{}\n{}".format(str(json_obj["id"]), json_obj["title"][3:])

    def tc_requirement(self, json_obj, row_idx, tc):
        return str(json_obj["custom_io_requirement"])

    def tc_expected_result(self, json_obj, row_idx, tc):
        return self.remove_picture_placeholder(self.get_expected_result(json_obj["custom_steps_separated"]))

    def tc_test_method(self, json_obj, row_idx, tc):
        return "{} / {}".format(json_obj["custom_string_objective_evidence"], json_obj["custom_test_objev"])

    def tr_step_num(self, json_obj, row_idx, tc):
        self.test_run_step_counter += 1
        return self.test_run_step_counter

    def tr_result_description(self, json_obj, row_idx, tc):
        test_result = []
        test_result.append("Test case: C{} - {}".format(tc, self.tc_js["title"][3:]))
        test_result.append("Test Result ID: T{}\n".format(json_obj["test_id"]))
        test_result.append(self.extract_all_step_result(json_obj["custom_step_results"]))
        return "\n".join(test_result)



    def tr_result(self, json_obj, row_idx, tc):
        return self.status_text[json_obj["status_id"]]

    def tr_date(self, json_obj, row_idx, tc):
        date = datetime.datetime.fromtimestamp(json_obj["created_on"]).strftime("%Y-%m-%d")
        return date

    def tr_initial(self, json_obj, row_idx, tc):
        try:
            initial = self.user_initial[json_obj["created_by"]]

        except KeyError:
            initial = "Not found id {}".format(json_obj["created_by"])
        return initial

    def change_table_font(self, table, font_size_pt):
        """Table: docx x table ex: d.table[0] where d in docx object"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(font_size_pt)

    def get_last_valid_result(self, json_object):
        for js_obj in json_object:
            if isinstance(js_obj["status_id"], int):
                return [js_obj]
        return 0

    def my_http_request(self, url, user, password):
        authentication = (user, password)
        head = {"Content-Type": "application/json"}
        return requests.get(url, auth=authentication, headers=head).json()

    def get_tc_list_from_section(self, tr_section_id):
        tc_list = []
        URL = "https://testrail.dwos.com//index.php?/api/v2/get_cases/" + str(self.config["project id"]) + "&section_id="
        authentication = (self.USER, self.PASSWORD)
        head = {"Content-Type": "application/json"}
        r = requests.get(URL + str(tr_section_id), auth=authentication, headers=head)
        for tc in range(len(r.json())):
            tc_list.append(r.json()[tc]["id"])
        return tc_list

def main():
    ex = Export_test_rail()
    user_initial = ex.create_user_initial()


    if ex.config["test report"]:   # test report creation
        print("Process test Report")
        for tr_section_id, table_id in ex.config["table mapping"].items():
            tc_list = ex.get_tc_list_from_section(tr_section_id)
            ex.test_run_step_counter = 0
            df_orginal = None
            for test_case_id in tc_list:
                url = "https://testrail.dwos.com/index.php?/api/v2/get_case/" + str(test_case_id)
                ex.tc_js = ex.my_http_request(url, ex.USER, ex.PASSWORD)

                url = "https://testrail.dwos.com//index.php?/api/v2/get_results_for_case/" + \
                      str(ex.config["test report run id"]) + "/" + str(test_case_id)
                ex.js = ex.my_http_request(url, ex.USER, ex.PASSWORD)
                try:
                    ex.js[0]
                    js = ex.get_last_valid_result(ex.js)
                    df = ex.create_df_from_json(js, ["Step", "Result Description", "Result", "Date", "Initial"],
                                             [ex.tr_step_num, ex.tr_result_description, ex.tr_result, ex.tr_date, ex.tr_initial], test_case_id)
                    df_orginal = pd.concat([df_orginal, df], ignore_index=True)
                except KeyError:
                    continue
        ex.write_to_doc_table(df_orginal, table_id, ex.config["template path"], ex.config["output doc name"])
    else:   # test spec creation
        for tr_section_id, table_id in ex.config["table mapping"].items():
            url = "https://testrail.dwos.com//index.php?/api/v2/get_cases/" + str(ex.config["project id"]) + "&section_id=" + str(tr_section_id)
            js = ex.my_http_request(url, ex.USER, ex.PASSWORD)
            df = ex.create_df_from_json(js, ["Step", "Description", "Requirement", "Expected Result",
                                         "Test Method / Objective Evidence"],
                                 [ex.step_num, ex.tc_description, ex.tc_requirement, ex.tc_expected_result, ex.tc_test_method])
            ex.write_to_doc_table(df, table_id, ex.config["template path"], ex.config["output doc name"])


main()

